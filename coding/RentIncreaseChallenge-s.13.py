import xlwings as xw
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# =============Input=============================
doc_filename = 'RentIncreaseChallenge-s.13.docx'
para1_body = '{Date}\n{Taddress}\n{TPostcode}'
para2_body = 'Dear {LLname},\n\nRE: Rent Increase at {Taddress}\n\nI am the tenant at the above address and I am writing to inform you that the recent/proposed is invalid for the following reason(s):'
b1_body = '{Q1} Landlords may only increase rent through a s.13 notice if the tenancy agreement contains a rent review clause. Because my tenancy agreement does not feature a rent review clause, the rent increase is invalid.'
b2_body = '{Q2} In order to increase rent, landlords must use \'Tenancy Form 4\' or a document with all the same information. Because of this, the rent increase in invalid.'
b3_body = '{Q3} Unless otherwise stated, landlords may only increase rent once per year. After reviewing my tenancy agreement, I have found there to be no term stating that you may increase my rent more than once a year. If I pay rent on a weekly or monthly basis, you must provide a minimum notice of one month. If I pay on a yearly basis, you must provide a minimum notice of one month. If I pay on a yearly basis, you must give six months.'
b4_body = '{Q4} Landlords must give tenants one month notice when increasing rent. Because of this, the rent increase is invalid.'
b5_body = '{Q5} An increase in rent may only be effective after the term has ended. Because of this, the rent increase is invalid.'
para3_body = '\n\nPlease contact me as soon as possible to further discuss the matter.\n\nBest regards,\n\n\n{Name}'

# =============Dir================================
excel_dir = '../data/Tenantchat.xlsx'

# ###############MAIN#########################################
def main():
    wb = xw.Book(excel_dir)
    sht = wb.sheets['RentIncrease-S.13']

    # -----------------------------------------------
    # Define variables 
    date = sht.range('A2').value
    taddress = sht.range('F2').value
    tpostcode = sht.range('G2').value
    llname = sht.range('E2').value
    name = sht.range('B2').value
    q1 = sht.range('H2').value
    q2 = sht.range('I2').value
    q3 = sht.range('J2').value
    q4 = sht.range('K2').value
    q5 = sht.range('L2').value

    # =================================================
    # Create document from here
    d = Document()

    # -------------------------------------------------
    para1 = d.add_paragraph(para1_body.format(
                                            Date= date,
                                            Taddress= taddress,
                                            TPostcode= tpostcode))

    para1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # -------------------------------------------------
    para2 = d.add_paragraph(para2_body.format(
                                            LLname= llname,
                                            Taddress= taddress))

    para2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # -------------------------------------------------
    if q1 == 'Yes':
        # b1 = d.add_paragraph('')    # empty with 3 lines as per the preview img
        pass
    elif q1 == 'No':
        b1 = d.add_paragraph(b1_body.format(Q1= ''), style= 'List Bullet')
    else:
        b1 = d.add_paragraph(b1_body.format(Q1= q1), style= 'List Bullet')
    # -------------------------------------------------
    if q2 == 'Yes':
        # b2 = d.add_paragraph('')    # empty with 2 lines as per the preview img
        pass
    elif q2 == 'No':
        b2 = d.add_paragraph(b2_body.format(Q2= ''), style= 'List Bullet')
    else:
        b2 = d.add_paragraph(b2_body.format(Q2= q2), style= 'List Bullet')
    # -------------------------------------------------
    if q3 == 'Yes':
        b3 = d.add_paragraph(b3_body.format(Q3= ''), style= 'List Bullet')
    elif q3 == 'No':
        # b3 = d.add_paragraph('')    # empty with 5 lines as per the preview img
        pass
    else:
        b3 = d.add_paragraph(b3_body.format(Q3= q3), style= 'List Bullet')
    # -------------------------------------------------
    if q4 == 'Yes':
        # b4 = d.add_paragraph('')    # empty with 5 lines as per the preview img
        pass
    elif q4 == 'No':
        b4 = d.add_paragraph(b4_body.format(Q4= ''), style= 'List Bullet')
    else:
        b4 = d.add_paragraph(b4_body.format(Q4= q4), style= 'List Bullet')
    # -------------------------------------------------
    if q5 == 'Yes':
        b5 = d.add_paragraph(b5_body.format(Q5= ''), style= 'List Bullet')
    elif q5 == 'No':
        # b5 = d.add_paragraph('\n\n')    # empty with 5 lines as per the preview img
        pass
    else:
        b5 = d.add_paragraph(b5_body.format(Q5= q5), style= 'List Bullet')
    # -------------------------------------------------
    para3 = d.add_paragraph(para3_body.format(Name= name))

    para3.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # -------------------------------------------------
    d.save(doc_filename)

#--------------------------------------------------------------------------------------------------------------------------------------------------------------------
# MAIN Function call
#--------------------------------------------------------------------------------------------------------------------------------------------------------------------
if __name__ == "__main__":
    main()
